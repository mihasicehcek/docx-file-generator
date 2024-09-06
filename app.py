from docx import Document
from docx.shared import Inches
from PIL import Image
import numpy as np
import os
from tqdm import tqdm
# import pypandoc

def generate_png_with_size(target_size_bytes, output_path='output_image.png'):
    # Начальные размеры изображения
    width, height = 100, 100

    # Генерируем изображение до тех пор, пока размер файла не приблизится к нужному
    while True:
        # Создание массива с случайными цветами
        data = np.random.randint(0, 256, (height, width, 3), dtype=np.uint8)

        # Создание изображения из массива
        image = Image.fromarray(data, 'RGB')

        # Сохранение изображения в файл
        image.save(output_path, 'PNG')

        # Проверка размера файла
        file_size = os.path.getsize(output_path)

        if file_size >= target_size_bytes:
            break

        # Увеличение размеров изображения
        width += 50
        height += 50


total_size_mb = float(input("Result file size in Mb: "))
num_pages = float(input("Number of pages: "))

total_size_in_bytes = total_size_mb * 1024 * 1024

size_per_page = total_size_in_bytes / num_pages

doc = Document()

for i in tqdm(range(1, int(num_pages) + 1), desc="Document creation", unit=" page"):
    if size_per_page > 0:
        generate_png_with_size(size_per_page, 'output_image.png')
        doc.add_picture('output_image.png', width=Inches(4))
        doc.add_page_break()

        doc.save('output_document.docx')
        current_size = os.path.getsize('output_document.docx')
        remain_size = total_size_in_bytes - current_size
        if (num_pages - i) != 0:
            size_per_page = remain_size / (num_pages - i)
    else:
        doc.add_paragraph('page ' + str(i) + ' of ' + str(num_pages) + ' pages')
        doc.add_page_break()

os.remove('output_image.png')

doc.save('output_document.docx')
# pypandoc.convert_file('output_document.docx', 'pdf', outputfile='output_document.pdf')

print("Result document size is "+str(int(os.path.getsize('output_document.docx')/1024/1024))+" Mb")